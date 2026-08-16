from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "competition"
PDF_PATH = OUT / "MoonResilience项目申报书.pdf"
DOCX_PATH = OUT / "MoonResilience项目申报书.docx"

TITLE = "MoonResilience 项目申报书"
SUBTITLE = "MoonBit 弹性治理基础库"
GITLINK = "https://www.gitlink.org.cn/q2weasd/MoonResilience"
GITHUB = "https://github.com/liying-han/MoonResilience"

META = [
    ("项目名称", "MoonResilience：MoonBit 弹性治理基础库"),
    ("参赛方向", "MoonBit 国产基础软件开源生态项目"),
    ("开源许可证", "Apache-2.0"),
    ("GitLink 仓库", GITLINK),
    ("GitHub 仓库", GITHUB),
]

SECTIONS = [
    (
        "01  项目目标与适用对象",
        "MoonResilience 面向使用 MoonBit 开发服务端程序、任务执行器、消息消费者和基础工具的开发者。项目集中处理调用链中反复出现的失败重试、故障隔离、流量控制和状态观察问题。核心库不绑定 HTTP 框架或系统时钟，调用方提供动作、策略状态和当前时间后，可在不同运行环境复用相同治理逻辑。",
    ),
    (
        "02  项目方案与核心能力",
        "Retry 支持固定、线性、指数和序列退避，并按最大次数、时间预算、错误码及可重试标记停止；Circuit Breaker 实现 Closed、Open、HalfOpen 状态机及恢复探测限制；Rate Limiter 提供 token bucket 和 fixed window；Bulkhead 管理并发槽位、有限等待队列、取消和 FIFO 晋升。统一执行链采用“限流、隔离、熔断、动作与重试”的固定顺序，返回最终状态、拒绝原因、尝试次数和执行轨迹。",
    ),
    (
        "03  技术路线",
        "项目使用纯 MoonBit 实现，不引入第三方依赖。策略对象均为显式状态值，函数返回更新后的新状态，不使用隐藏全局变量。时间相关策略统一使用毫秒整数，测试与演示通过 VirtualClock 推进时间。Config 负责配置解析与约束诊断；Telemetry 记录事件、计数和延迟分桶并导出 Prometheus 文本；Diagnostics 检查健康状态和运行时不变量；Simulation 使用预设响应复现故障与恢复过程。",
    ),
    (
        "04  当前成果与验证方式",
        "仓库已完成核心库、CLI 示例、配置样例、API 与架构文档，现有约 5.0k 行有效 MoonBit 代码和 110 项测试。测试覆盖退避边界、熔断状态转换、两类限流器、Bulkhead 容量耗尽、策略顺序、配置错误、时间模拟、指标导出及批量场景。评审者可直接运行 moon check、moon test、moon build 和 moon run cmd/main；README 同时列出同步执行、显式时钟等当前边界。",
    ),
    (
        "05  公开开发与后续计划",
        "GitLink 与 GitHub 保存相同版本，提交按模块、测试、文档和发布材料拆分。后续工作通过公开工单和功能分支记录，优先完成 HTTP 调用适配、等待队列超时、滑动窗口限流、事件订阅和跨后端一致性验证。阶段交付以测试可运行、接口文档同步和变更记录完整为验收条件。",
    ),
]


def register_pdf_fonts():
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    bold_path = Path(r"C:\Windows\Fonts\msyhbd.ttc")
    if not font_path.exists():
        raise FileNotFoundError("Microsoft YaHei font was not found")
    pdfmetrics.registerFont(TTFont("YaHei", str(font_path)))
    pdfmetrics.registerFont(TTFont("YaHeiBold", str(bold_path if bold_path.exists() else font_path)))


def build_pdf():
    register_pdf_fonts()
    forest = colors.HexColor("#214E3B")
    green = colors.HexColor("#2F6B57")
    gold = colors.HexColor("#A77A35")
    pale = colors.HexColor("#EAF2ED")
    warm_pale = colors.HexColor("#F5F0E7")
    rule = colors.HexColor("#AFC2B8")
    ink = colors.HexColor("#202A26")
    muted = colors.HexColor("#5C6963")

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=11 * mm,
        bottomMargin=10 * mm,
        title=TITLE,
        author="MoonResilience",
    )
    styles = {
        "kicker": ParagraphStyle(
            "Kicker", fontName="YaHeiBold", fontSize=8.2, leading=10, textColor=gold,
            alignment=TA_LEFT, spaceAfter=2,
        ),
        "title": ParagraphStyle(
            "Title", fontName="YaHeiBold", fontSize=20.5, leading=24, textColor=forest,
            alignment=TA_LEFT, spaceAfter=2,
        ),
        "subtitle": ParagraphStyle(
            "Subtitle", fontName="YaHei", fontSize=9.8, leading=13, textColor=muted,
            alignment=TA_LEFT, spaceAfter=7,
        ),
        "label": ParagraphStyle(
            "Label", fontName="YaHeiBold", fontSize=8.2, leading=10.2, textColor=forest,
        ),
        "meta": ParagraphStyle(
            "Meta", fontName="YaHei", fontSize=8.0, leading=10.0, textColor=ink,
        ),
        "heading": ParagraphStyle(
            "Heading", fontName="YaHeiBold", fontSize=9.6, leading=12, textColor=forest,
            spaceBefore=0, spaceAfter=2.2,
        ),
        "body": ParagraphStyle(
            "Body", fontName="YaHei", fontSize=8.9, leading=15.1, textColor=ink,
            alignment=TA_LEFT, firstLineIndent=0, spaceAfter=0,
        ),
        "footer": ParagraphStyle(
            "Footer", fontName="YaHei", fontSize=7.4, leading=9, textColor=muted,
            alignment=TA_CENTER,
        ),
        "note": ParagraphStyle(
            "Note", fontName="YaHei", fontSize=7.8, leading=10.5, textColor=muted,
            alignment=TA_LEFT,
        ),
    }

    masthead_rule = Table([["", ""]], colWidths=[138 * mm, 42 * mm], rowHeights=[2.2 * mm])
    masthead_rule.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), green),
        ("BACKGROUND", (1, 0), (1, 0), gold),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    story = [
        masthead_rule,
        Spacer(1, 5),
        Paragraph("2026 MoonBit 国产基础软件开源大赛", styles["kicker"]),
        Paragraph(TITLE, styles["title"]),
        Paragraph(SUBTITLE + " · 项目标识 moonresilience", styles["subtitle"]),
    ]
    meta_rows = []
    for label, value in META:
        shown = value
        if label in ("GitLink 仓库", "GitHub 仓库"):
            shown = "<link href='" + value + "' color='#2F6B57'>" + value + "</link>"
        meta_rows.append([Paragraph(label, styles["label"]), Paragraph(shown, styles["meta"])])
    table = Table(meta_rows, colWidths=[31 * mm, 149 * mm], hAlign="CENTER")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), warm_pale),
        ("BOX", (0, 0), (-1, -1), 0.55, rule),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, rule),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 2.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
    ]))
    story.extend([table, Spacer(1, 6)])
    for index, (heading, body) in enumerate(SECTIONS):
        block = Table(
            [[Paragraph(heading, styles["heading"])], [Paragraph(body, styles["body"])]],
            colWidths=[180 * mm],
        )
        block.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, 0), pale),
            ("LINEBEFORE", (0, 0), (0, -1), 2.1, green),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, 0), 3.1),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 2.1),
            ("TOPPADDING", (0, 1), (-1, 1), 4.5),
            ("BOTTOMPADDING", (0, 1), (-1, 1), 5.0),
        ]))
        story.append(KeepTogether(block))
        if index != len(SECTIONS) - 1:
            story.append(Spacer(1, 5.0))
    metric_cells = [
        ("4 类", "基础策略"),
        ("2 种", "限流算法"),
        ("约 5.0k 行", "有效 MoonBit 代码"),
        ("110 项", "自动化测试"),
    ]
    metric_table = Table(
        [[
            Paragraph(
                "<b>" + value + "</b><br/><font size='7.2'>" + label + "</font>",
                ParagraphStyle(
                    "Metric", fontName="YaHei", fontSize=9.1, leading=11.5,
                    textColor=forest, alignment=TA_CENTER,
                ),
            )
            for value, label in metric_cells
        ]],
        colWidths=[45 * mm] * 4,
    )
    metric_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), warm_pale),
        ("BOX", (0, 0), (-1, -1), 0.5, rule),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, rule),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.extend([
        Spacer(1, 7),
        metric_table,
        Spacer(1, 5),
        Table(
            [[Paragraph("阶段状态：基础原型、CLI、测试和接口文档已完成，可直接复现主要策略行为。", styles["note"])],
             [Paragraph("当前边界：核心库不直接发起网络请求；生产接入层负责提供实际时钟与异步调度。", styles["note"])]],
            colWidths=[180 * mm],
            style=TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F1F5F2")),
                ("LINEBEFORE", (0, 0), (0, -1), 1.5, rule),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, 0), 4),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 2),
                ("TOPPADDING", (0, 1), (-1, 1), 2),
                ("BOTTOMPADDING", (0, 1), (-1, 1), 4),
            ]),
        ),
        Spacer(1, 6),
        Paragraph(
            "验证命令：moon info · moon fmt · moon check --warn-list +73 · moon test · moon run cmd/main",
            styles["footer"],
        ),
    ])
    doc.build(story)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=45, start=90, bottom=45, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size, bold=False, color="17212B"):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(8.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.18

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(1)
    set_run_font(p.add_run("2026 MoonBit 国产基础软件开源大赛"), 8.2, color="A77A35")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(TITLE), 20, bold=True, color="214E3B")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run(SUBTITLE + " · 项目标识 moonresilience"), 9.5, color="526170")

    table = doc.add_table(rows=len(META), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(3.1)
    table.columns[1].width = Cm(14.9)
    for row, (label, value) in zip(table.rows, META):
        row.cells[0].width = Cm(3.1)
        row.cells[1].width = Cm(14.9)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        set_cell_shading(row.cells[0], "F5F0E7")
        left = row.cells[0].paragraphs[0]
        right = row.cells[1].paragraphs[0]
        set_run_font(left.add_run(label), 8.0, bold=True, color="214E3B")
        set_run_font(right.add_run(value), 7.8)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    for heading, body in SECTIONS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2.5)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.keep_with_next = True
        set_run_font(p.add_run(heading), 9.4, bold=True, color="214E3B")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.34
        set_run_font(p.add_run(body), 8.7)

    metrics = doc.add_table(rows=1, cols=4)
    metric_cells = [
        ("4 类", "基础策略"),
        ("2 种", "限流算法"),
        ("约 5.0k 行", "有效 MoonBit 代码"),
        ("110 项", "自动化测试"),
    ]
    for cell, (value, label) in zip(metrics.rows[0].cells, metric_cells):
        set_cell_shading(cell, "F5F0E7")
        set_cell_margins(cell, top=70, bottom=70)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(value + "\n"), 9.0, bold=True, color="214E3B")
        set_run_font(paragraph.add_run(label), 7.2, color="526170")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    set_run_font(
        p.add_run("阶段状态：基础原型、CLI、测试和接口文档已完成，可直接复现主要策略行为。"),
        7.6,
        color="526170",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    set_run_font(
        p.add_run("当前边界：核心库不直接发起网络请求；生产接入层负责提供实际时钟与异步调度。"),
        7.6,
        color="526170",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    set_run_font(
        p.add_run("验证命令：moon info · moon fmt · moon check --warn-list +73 · moon test · moon run cmd/main"),
        7.2,
        color="526170",
    )
    doc.core_properties.title = TITLE
    doc.core_properties.subject = SUBTITLE
    doc.core_properties.author = "MoonResilience"
    doc.core_properties.keywords = "MoonBit, resilience, OSC2026"
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build_pdf()
    build_docx()
    print(PDF_PATH)
    print(DOCX_PATH)
